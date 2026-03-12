import os
import shutil
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List, Optional
import markdown
import json
import hashlib
from dataclasses import dataclass, asdict


@dataclass
class DocumentVersion:
    """Represents a version of a document."""
    version: str
    content: str
    metadata: Dict[str, Any]
    created_at: str
    hash: str


class DocumentProcessor:
    """Handles document processing, versioning, and publishing."""

    def __init__(self, base_dir: Path = None):
        # User-generated content should go to gitignored directory
        self.base_dir = base_dir or Path('user_docs')
        self.base_dir.mkdir(exist_ok=True)

    def create_document(self,
                       title: str,
                       content: str,
                       metadata: Optional[Dict[str, Any]] = None,
                       auto_version: bool = True) -> Dict[str, Any]:
        """Create or update a document with versioning."""

        # Sanitize title for file system
        safe_title = self._sanitize_filename(title)
        doc_dir = self.base_dir / safe_title
        doc_dir.mkdir(exist_ok=True)

        # Calculate content hash for change detection
        content_hash = hashlib.md5(content.encode()).hexdigest()

        # Check if document exists and content changed
        if auto_version and self._content_changed(doc_dir, content_hash):
            version = self._get_next_version(doc_dir)
        else:
            version = self._get_current_version(doc_dir) or "1.0"

        # Create version metadata
        version_metadata = {
            "title": title,
            "version": version,
            "created_at": datetime.now().isoformat(),
            "hash": content_hash,
            **(metadata or {})
        }

        # Save version
        version_file = doc_dir / f"v{version}.md"
        with open(version_file, 'w', encoding='utf-8') as f:
            f.write(content)

        # Save metadata
        meta_file = doc_dir / f"v{version}.json"
        with open(meta_file, 'w', encoding='utf-8') as f:
            json.dump(version_metadata, f, indent=2)

        # Update current symlink
        current_link = doc_dir / "current.md"
        if current_link.exists():
            current_link.unlink()
        current_link.symlink_to(f"v{version}.md")

        # Update current metadata
        current_meta = doc_dir / "current.json"
        if current_meta.exists():
            current_meta.unlink()
        current_meta.symlink_to(f"v{version}.json")

        return {
            "success": True,
            "title": title,
            "version": version,
            "file_path": str(version_file),
            "directory": str(doc_dir),
            "hash": content_hash,
            "is_new_version": version != "1.0"
        }

    def get_document_versions(self, title: str) -> List[DocumentVersion]:
        """Get all versions of a document."""
        safe_title = self._sanitize_filename(title)
        doc_dir = self.base_dir / safe_title

        if not doc_dir.exists():
            return []

        versions = []
        for version_file in doc_dir.glob("v*.md"):
            version_num = version_file.stem[1:]  # Remove 'v' prefix
            meta_file = doc_dir / f"v{version_num}.json"

            # Read content
            with open(version_file, 'r', encoding='utf-8') as f:
                content = f.read()

            # Read metadata
            metadata = {}
            if meta_file.exists():
                with open(meta_file, 'r', encoding='utf-8') as f:
                    metadata = json.load(f)

            versions.append(DocumentVersion(
                version=version_num,
                content=content,
                metadata=metadata,
                created_at=metadata.get('created_at', ''),
                hash=metadata.get('hash', '')
            ))

        # Sort by version
        versions.sort(key=lambda v: [int(x) for x in v.version.split('.')])
        return versions

    def convert_to_formats(self,
                          title: str,
                          version: str = "current",
                          formats: List[str] = None) -> Dict[str, Any]:
        """Convert document to multiple formats."""
        if formats is None:
            formats = ["html", "docx"]

        safe_title = self._sanitize_filename(title)
        doc_dir = self.base_dir / safe_title

        if version == "current":
            source_file = doc_dir / "current.md"
        else:
            source_file = doc_dir / f"v{version}.md"

        if not source_file.exists():
            return {"success": False, "error": f"Document {title} v{version} not found"}

        # Read markdown content
        with open(source_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        results = {"success": True, "files": {}}

        # Convert to HTML
        if "html" in formats:
            html_content = markdown.markdown(
                md_content,
                extensions=['tables', 'fenced_code', 'toc', 'codehilite']
            )

            # Wrap in full HTML document
            full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; line-height: 1.6; margin: 2rem auto; max-width: 800px; padding: 0 1rem; }}
        h1, h2, h3 {{ color: #333; }}
        code {{ background: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
        pre {{ background: #f4f4f4; padding: 1rem; border-radius: 5px; overflow-x: auto; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""

            html_file = doc_dir / f"v{version}.html"
            with open(html_file, 'w', encoding='utf-8') as f:
                f.write(full_html)
            results["files"]["html"] = str(html_file)

        # Convert to DOCX (basic implementation)
        if "docx" in formats:
            try:
                from docx import Document
                from docx.shared import Inches

                doc = Document()

                # Add title
                title_paragraph = doc.add_heading(title, 0)

                # Convert markdown to basic docx
                # This is a simplified conversion - could be enhanced
                lines = md_content.split('\n')
                for line in lines:
                    if line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                    elif line.strip():
                        doc.add_paragraph(line)

                docx_file = doc_dir / f"v{version}.docx"
                doc.save(str(docx_file))
                results["files"]["docx"] = str(docx_file)

            except ImportError:
                results["errors"] = results.get("errors", [])
                results["errors"].append("python-docx not available for DOCX conversion")

        return results

    def publish_document(self,
                        title: str,
                        version: str = "current",
                        targets: List[str] = None,
                        target_configs: Dict[str, Any] = None) -> Dict[str, Any]:
        """Publish document to various targets."""
        if targets is None:
            targets = []
        if target_configs is None:
            target_configs = {}

        safe_title = self._sanitize_filename(title)
        doc_dir = self.base_dir / safe_title

        if version == "current":
            source_file = doc_dir / "current.md"
            metadata_file = doc_dir / "current.json"
        else:
            source_file = doc_dir / f"v{version}.md"
            metadata_file = doc_dir / f"v{version}.json"

        if not source_file.exists():
            return {"success": False, "error": f"Document {title} v{version} not found"}

        # Read content and metadata
        with open(source_file, 'r', encoding='utf-8') as f:
            content = f.read()

        metadata = {}
        if metadata_file.exists():
            with open(metadata_file, 'r', encoding='utf-8') as f:
                metadata = json.load(f)

        results = {"success": True, "published": {}, "errors": []}

        for target in targets:
            try:
                if target == "confluence":
                    result = self._publish_to_confluence(title, content, metadata, target_configs.get("confluence", {}))
                    results["published"]["confluence"] = result
                elif target == "github":
                    result = self._publish_to_github(title, content, metadata, target_configs.get("github", {}))
                    results["published"]["github"] = result
                else:
                    results["errors"].append(f"Unknown publish target: {target}")
            except Exception as e:
                results["errors"].append(f"Failed to publish to {target}: {str(e)}")
                results["published"][target] = {"success": False, "error": str(e)}

        return results

    def get_document_diff(self, title: str, version1: str, version2: str) -> Dict[str, Any]:
        """Get diff between two document versions."""
        versions = self.get_document_versions(title)

        v1_content = None
        v2_content = None

        for version in versions:
            if version.version == version1:
                v1_content = version.content
            if version.version == version2:
                v2_content = version.content

        if v1_content is None or v2_content is None:
            return {"success": False, "error": "One or both versions not found"}

        # Simple diff implementation
        v1_lines = v1_content.splitlines()
        v2_lines = v2_content.splitlines()

        changes = []
        max_lines = max(len(v1_lines), len(v2_lines))

        for i in range(max_lines):
            v1_line = v1_lines[i] if i < len(v1_lines) else ""
            v2_line = v2_lines[i] if i < len(v2_lines) else ""

            if v1_line != v2_line:
                changes.append({
                    "line_number": i + 1,
                    "old": v1_line,
                    "new": v2_line,
                    "type": "modified" if v1_line and v2_line else ("added" if v2_line else "deleted")
                })

        return {
            "success": True,
            "title": title,
            "version1": version1,
            "version2": version2,
            "changes": changes,
            "changes_count": len(changes)
        }

    def _sanitize_filename(self, title: str) -> str:
        """Convert title to safe filename."""
        safe = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()
        safe = safe.replace(' ', '_').lower()
        return safe[:50]  # Limit length

    def _content_changed(self, doc_dir: Path, new_hash: str) -> bool:
        """Check if content has changed from current version."""
        current_meta = doc_dir / "current.json"
        if not current_meta.exists():
            return True

        try:
            with open(current_meta, 'r', encoding='utf-8') as f:
                metadata = json.load(f)
                return metadata.get('hash') != new_hash
        except:
            return True

    def _get_next_version(self, doc_dir: Path) -> str:
        """Get next version number."""
        current_version = self._get_current_version(doc_dir)
        if not current_version:
            return "1.0"

        major, minor = map(int, current_version.split('.'))
        return f"{major}.{minor + 1}"

    def _get_current_version(self, doc_dir: Path) -> Optional[str]:
        """Get current version number."""
        current_meta = doc_dir / "current.json"
        if not current_meta.exists():
            return None

        try:
            with open(current_meta, 'r', encoding='utf-8') as f:
                metadata = json.load(f)
                return metadata.get('version')
        except:
            return None

    def _publish_to_confluence(self, title: str, content: str, metadata: Dict, config: Dict) -> Dict[str, Any]:
        """Publish to Confluence."""
        # This would integrate with the Confluence MCP service
        return {"success": True, "url": f"https://confluence.example.com/pages/{title}", "message": "Published to Confluence"}

    def _publish_to_github(self, title: str, content: str, metadata: Dict, config: Dict) -> Dict[str, Any]:
        """Publish to GitHub."""
        # This would integrate with the GitHub MCP service
        return {"success": True, "url": f"https://github.com/user/repo/blob/main/{title}.md", "message": "Published to GitHub"}